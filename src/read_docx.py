from docx import Document
import os

def read_docx(file_path):
    try:
        # Проверяем существование файла
        if not os.path.exists(file_path):
            print(f"Ошибка: Файл {file_path} не найден")
            return None

        # Открываем документ
        doc = Document(file_path)
        
        # Читаем текст из всех параграфов
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Пропускаем пустые параграфы
                full_text.append(paragraph.text)
        
        # Читаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():  # Пропускаем пустые ячейки
                        row_text.append(cell.text)
                if row_text:  # Добавляем только непустые строки
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    except Exception as e:
        print(f"Ошибка при чтении файла: {str(e)}")
        return None

def analyze_document(text):
    if not text:
        return "Нет данных для анализа"
    
    # Базовый анализ документа
    lines = text.split('\n')
    analysis = {
        "Количество строк": len(lines),
        "Количество символов": len(text),
        "Количество слов": len(text.split()),
        "Содержит таблицы": " | " in text,
        "Содержит формулы": any(char in text for char in ['∫', '∑', '√', 'π', '∞'])
    }
    
    return analysis

def main():
    # Путь к файлу .docx
    file_path = "docs/Tsyrenov_VKR.docx"  # Замените на путь к вашему файлу
    
    # Читаем документ
    text = read_docx(file_path)
    
    if text:
        print("Содержимое документа:")
        print("-" * 50)
        print(text)
        print("-" * 50)
        
        # Анализируем документ
        print("\nАнализ документа:")
        print("-" * 50)
        analysis = analyze_document(text)
        for key, value in analysis.items():
            print(f"{key}: {value}")
    else:
        print("Не удалось прочитать документ")

if __name__ == "__main__":
    main() 